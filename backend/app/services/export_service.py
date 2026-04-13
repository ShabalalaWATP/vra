"""Report export — Professional PDF and DOCX generation."""

import io
import base64
import json
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.models.dependency import Dependency, DependencyFinding
from app.models.file import File
from app.models.finding import Finding
from app.models.report import ExportArtifact, Report
from app.models.secret_candidate import SecretCandidate
from app.services.report_diagrams import (
    RenderedReportDiagram,
    guess_diagram_media_type,
    is_svg_bytes,
    render_report_diagrams,
    svg_to_png_bytes,
)

MAX_CODE_LINES = 4
MAX_POC_LINES = 120
SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}


@dataclass
class ExportSecretRow:
    type: str
    confidence: float | None
    file_path: str | None


def _truncate_code(snippet: str | None, max_lines: int = MAX_CODE_LINES) -> str:
    if not snippet:
        return ""
    lines = snippet.strip().splitlines()
    if len(lines) <= max_lines:
        return snippet.strip()
    return "\n".join(lines[:max_lines]) + f"\n... ({len(lines) - max_lines} more lines)"


def _sort_findings(findings: list[Finding]) -> list[Finding]:
    return sorted(findings, key=lambda f: (SEVERITY_ORDER.get(f.severity, 9), -(f.confidence or 0)))


def _truncate_text(value: str | None, max_chars: int) -> str:
    if not value:
        return ""
    text = value.strip()
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3].rstrip() + "..."


def _text_blocks(value: str | None) -> list[str]:
    if not value:
        return []
    return [block.strip() for block in str(value).split("\n\n") if block.strip()]


def _clean_string_list(values) -> list[str]:
    if not isinstance(values, list):
        return []
    cleaned: list[str] = []
    for value in values:
        text = str(value).strip()
        if text:
            cleaned.append(text)
    return cleaned


def _normalise_exploit_evidence(finding: Finding) -> dict:
    payload = finding.exploit_evidence if isinstance(finding.exploit_evidence, dict) else {}
    normalised = {
        "difficulty": str(payload.get("difficulty") or finding.exploit_difficulty or "").strip() or None,
        "target_route": str(payload.get("target_route") or "").strip() or None,
        "prerequisites": _clean_string_list(payload.get("prerequisites") or finding.exploit_prerequisites),
        "validation_steps": _clean_string_list(payload.get("validation_steps")),
        "cleanup_notes": _clean_string_list(payload.get("cleanup_notes")),
        "exploit_template": str(payload.get("exploit_template") or finding.exploit_template or "").strip() or None,
        "attack_scenario": str(payload.get("attack_scenario") or finding.attack_scenario or "").strip() or None,
        "components": _clean_string_list(payload.get("components")),
        "related_entry_points": _clean_string_list(payload.get("related_entry_points")),
        "related_taint_flows": _clean_string_list(payload.get("related_taint_flows")),
    }
    return {
        key: value
        for key, value in normalised.items()
        if value not in (None, [], "")
    }


def _finding_has_exploit_evidence(finding: Finding) -> bool:
    evidence = _normalise_exploit_evidence(finding)
    return any(
        evidence.get(key)
        for key in (
            "exploit_template",
            "attack_scenario",
            "target_route",
            "validation_steps",
            "cleanup_notes",
            "prerequisites",
            "related_entry_points",
            "related_taint_flows",
            "difficulty",
        )
    )


def _add_docx_exploit_evidence(
    doc,
    finding: Finding,
    *,
    heading_level: int = 3,
    heading_text: str | None = "Exploit Evidence",
) -> None:
    from docx.shared import Pt, RGBColor

    evidence = _normalise_exploit_evidence(finding)
    if not evidence:
        return

    if heading_text:
        doc.add_heading(heading_text, level=heading_level)

    if evidence.get("difficulty"):
        p = doc.add_paragraph()
        p.add_run("Exploit Difficulty: ").bold = True
        p.add_run(str(evidence["difficulty"]).upper())

    if evidence.get("target_route"):
        p = doc.add_paragraph()
        p.add_run("Target Route / Invocation: ").bold = True
        p.add_run(str(evidence["target_route"]))

    if evidence.get("components"):
        p = doc.add_paragraph()
        p.add_run("Components: ").bold = True
        p.add_run(", ".join(evidence["components"][:6]))

    if evidence.get("related_entry_points"):
        p = doc.add_paragraph()
        p.add_run("Related Entry Points: ").bold = True
        p.add_run("; ".join(evidence["related_entry_points"][:4]))

    if evidence.get("prerequisites"):
        p = doc.add_paragraph()
        p.add_run("Prerequisites: ").bold = True
        p.add_run(", ".join(evidence["prerequisites"][:8]))

    if evidence.get("attack_scenario"):
        p = doc.add_paragraph()
        p.add_run("Attack Scenario: ").bold = True
        p.add_run(_truncate_text(str(evidence["attack_scenario"]), 1500))

    if evidence.get("related_taint_flows"):
        p = doc.add_paragraph()
        p.add_run("Related Taint Flows: ").bold = True
        p.add_run("; ".join(evidence["related_taint_flows"][:4]))

    if evidence.get("validation_steps"):
        p = doc.add_paragraph()
        p.add_run("Validation Steps").bold = True
        for step in evidence["validation_steps"][:6]:
            doc.add_paragraph(step, style="List Bullet")

    if evidence.get("cleanup_notes"):
        p = doc.add_paragraph()
        p.add_run("Cleanup Notes").bold = True
        for step in evidence["cleanup_notes"][:6]:
            doc.add_paragraph(step, style="List Bullet")

    if evidence.get("exploit_template"):
        p = doc.add_paragraph()
        p.add_run("Proof of Concept").bold = True
        code_p = doc.add_paragraph()
        run = code_p.add_run(_truncate_code(str(evidence["exploit_template"]), MAX_POC_LINES))
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)


def _render_html_exploit_evidence(
    finding: Finding,
    *,
    heading_text: str | None = "Exploit Evidence",
) -> list[str]:
    evidence = _normalise_exploit_evidence(finding)
    if not evidence:
        return []

    parts: list[str] = []
    if heading_text:
        parts.append(f"<h3>{_esc(heading_text)}</h3>")

    if evidence.get("difficulty"):
        parts.append(f"<p><strong>Exploit Difficulty:</strong> {_esc(str(evidence['difficulty']).upper())}</p>")
    if evidence.get("target_route"):
        parts.append(f"<p><strong>Target Route / Invocation:</strong> {_esc(str(evidence['target_route']))}</p>")
    if evidence.get("components"):
        parts.append(f"<p><strong>Components:</strong> {_esc(', '.join(evidence['components'][:6]))}</p>")
    if evidence.get("related_entry_points"):
        parts.append(f"<p><strong>Related Entry Points:</strong> {_esc('; '.join(evidence['related_entry_points'][:4]))}</p>")
    if evidence.get("prerequisites"):
        parts.append(f"<p><strong>Prerequisites:</strong> {_esc(', '.join(evidence['prerequisites'][:8]))}</p>")
    if evidence.get("attack_scenario"):
        parts.append(f"<p><strong>Attack Scenario:</strong> {_esc(_truncate_text(str(evidence['attack_scenario']), 1500))}</p>")
    if evidence.get("related_taint_flows"):
        parts.append(f"<p><strong>Related Taint Flows:</strong> {_esc('; '.join(evidence['related_taint_flows'][:4]))}</p>")
    if evidence.get("validation_steps"):
        parts.append("<p><strong>Validation Steps:</strong></p><ul>")
        for step in evidence["validation_steps"][:6]:
            parts.append(f"<li>{_esc(step)}</li>")
        parts.append("</ul>")
    if evidence.get("cleanup_notes"):
        parts.append("<p><strong>Cleanup Notes:</strong></p><ul>")
        for step in evidence["cleanup_notes"][:6]:
            parts.append(f"<li>{_esc(step)}</li>")
        parts.append("</ul>")
    if evidence.get("exploit_template"):
        parts.append("<p><strong>Proof of Concept:</strong></p>")
        parts.append(
            "<pre class='poc-block'><code>"
            f"{_esc(_truncate_code(str(evidence['exploit_template']), MAX_POC_LINES))}"
            "</code></pre>"
        )
    return parts


def _parse_report_architecture(report: Report) -> dict:
    raw = report.architecture
    if not raw:
        return {}
    if isinstance(raw, dict):
        return raw
    try:
        payload = json.loads(str(raw))
        return payload if isinstance(payload, dict) else {}
    except Exception:
        return {}


def _format_entry_point(entry_point) -> str:
    if isinstance(entry_point, dict):
        method = str(entry_point.get("method") or "").strip()
        path = str(entry_point.get("path") or "").strip()
        entry_type = str(entry_point.get("type") or "").strip()
        file_path = str(entry_point.get("file") or "").strip()
        function = str(entry_point.get("function") or "").strip()

        headline = " ".join(part for part in [method, path] if part).strip()
        if not headline:
            headline = entry_type or function or file_path or "entry point"

        location_bits = []
        if file_path:
            location_bits.append(file_path)
        if function:
            location_bits.append(function)
        if location_bits:
            headline += f" in {'::'.join(location_bits)}"
        if entry_type and entry_type not in headline:
            headline += f" [{entry_type}]"
        return headline
    return str(entry_point).strip()


def _collect_attack_surface_points(report: Report) -> list[str]:
    payload = _parse_report_architecture(report)
    points: list[str] = []

    for point in payload.get("attack_surface") or []:
        text = _format_entry_point(point)
        if text and text not in points:
            points.append(text)

    for entry_point in payload.get("entry_points") or []:
        text = _format_entry_point(entry_point)
        if text and text not in points:
            points.append(text)

    return points


def _collect_trust_boundaries(report: Report) -> list[str]:
    payload = _parse_report_architecture(report)
    return [str(item).strip() for item in (payload.get("trust_boundaries") or []) if str(item).strip()]


def _extract_exploit_chain_steps(finding: Finding) -> list[str]:
    steps: list[str] = []
    for evidence in getattr(finding, "evidence", []) or []:
        if getattr(evidence, "type", "") == "supporting":
            text = str(getattr(evidence, "description", "") or "").strip()
            if text:
                steps.append(text)

    if steps:
        return steps

    scenario = str(getattr(finding, "attack_scenario", "") or "").strip()
    if scenario:
        for line in scenario.splitlines():
            text = line.strip().lstrip("-").strip()
            if text:
                steps.append(text)
    return steps


def _scan_coverage_flags(cov: dict) -> list[str]:
    flags: list[str] = []
    if cov.get("degraded_coverage"):
        flags.append("Scanner coverage degraded")
    if cov.get("is_apk"):
        flags.append("APK decompilation")
    if cov.get("is_monorepo"):
        flags.append("Monorepo")
    if cov.get("obfuscated_files", 0) > 0:
        flags.append(f"{cov.get('obfuscated_files', 0)} obfuscated files")
    if cov.get("files_skipped_size", 0) > 0:
        flags.append(f"{cov.get('files_skipped_size', 0)} skipped due to size limit")
    if cov.get("files_skipped_cap", 0) > 0:
        flags.append(f"{cov.get('files_skipped_cap', 0)} skipped by scan cap")
    if (cov.get("ignored_file_count") or 0) > 0:
        flags.append(f"{cov.get('ignored_file_count', 0)} ignored by scope policy")
    if cov.get("has_doc_intelligence"):
        flags.append(f"{cov.get('doc_files_read', 0)} docs analysed")
    return flags


def _scanner_run_rows(cov: dict) -> list[dict]:
    rows = list((cov.get("scanner_runs") or {}).values())
    status_order = {"failed": 0, "degraded": 1, "completed": 2, "skipped": 3}
    rows.sort(key=lambda row: (status_order.get(str(row.get("status")), 9), str(row.get("scanner", "")).lower()))
    return rows


def _diagram_to_data_uri(diagram: RenderedReportDiagram) -> str | None:
    if not diagram.image_bytes:
        return None
    media_type = diagram.media_type or guess_diagram_media_type(diagram.image_bytes)
    if not media_type:
        return None
    encoded = base64.b64encode(diagram.image_bytes).decode("ascii")
    return f"data:{media_type};base64,{encoded}"


def _sort_dependency_findings(dep_findings: list[tuple[DependencyFinding, Dependency]]) -> list[tuple[DependencyFinding, Dependency]]:
    return sorted(
        dep_findings,
        key=lambda item: (
            item[0].risk_score is None,
            -(item[0].risk_score or 0.0),
            SEVERITY_ORDER.get((item[0].severity or "info").lower(), 9),
            item[1].name.lower(),
        ),
    )


def _humanise_dependency_label(value: str | None) -> str:
    if not value:
        return ""
    return value.replace("_", " ").strip().title()


def _format_dependency_risk_score(score: float | None) -> str:
    return f"{round(score)}/1000" if score is not None else "n/a"


def _format_dependency_exposure(df: DependencyFinding) -> str:
    reachability = _humanise_dependency_label(df.reachability_status)
    if reachability and df.reachability_confidence is not None:
        reachability = f"{reachability} ({round(df.reachability_confidence * 100)}%)"

    parts = [
        _humanise_dependency_label(df.relevance),
        reachability,
        _humanise_dependency_label(df.evidence_type),
    ]
    return " | ".join(part for part in parts if part)


def _summarise_dependency_usage(usage_evidence: list | None, limit: int = 2) -> str:
    if not usage_evidence:
        return ""

    kind_labels = {
        "import": "import",
        "reference": "reference",
        "vulnerable_function": "vulnerable function",
    }
    parts: list[str] = []
    for hit in usage_evidence[:limit]:
        if not isinstance(hit, dict):
            continue
        label = kind_labels.get(str(hit.get("kind", "")), _humanise_dependency_label(str(hit.get("kind", "usage"))).lower())
        symbol = str(hit.get("symbol", "")).strip()
        location = str(hit.get("file", "")).strip()
        line = hit.get("line")
        if location and isinstance(line, int):
            location = f"{location}:{line}"
        summary = label
        if symbol:
            summary += f" {symbol}"
        if location:
            summary += f" in {location}"
        parts.append(summary)

    return "; ".join(part for part in parts if part)


def _summarise_dependency_risk_factors(risk_factors: dict | None, limit: int = 3) -> str:
    if not isinstance(risk_factors, dict):
        return ""

    labels = {
        "reachability": "reachability",
        "relevance": "package usage",
        "vulnerable_function_match": "function match",
        "dev_dependency": "dev-only scope",
        "fix_available": "fix available",
        "hot_file_usage": "hot-file usage",
    }
    items = [
        (key, value)
        for key, value in risk_factors.items()
        if key not in {"base", "final"} and isinstance(value, (int, float)) and value != 0
    ]
    items.sort(key=lambda item: abs(float(item[1])), reverse=True)
    return ", ".join(
        f"{labels.get(key, _humanise_dependency_label(key).lower())} {'+' if value > 0 else ''}{round(float(value))}"
        for key, value in items[:limit]
    )


def _format_dependency_notes(df: DependencyFinding) -> str:
    notes: list[str] = []
    if df.ai_assessment:
        notes.append(_truncate_text(df.ai_assessment, 180))
    elif df.summary:
        notes.append(_truncate_text(df.summary, 180))

    usage = _summarise_dependency_usage(df.usage_evidence)
    if usage:
        notes.append(f"Usage: {usage}")

    if df.vulnerable_functions:
        notes.append(f"Functions: {', '.join(df.vulnerable_functions[:3])}")

    factors = _summarise_dependency_risk_factors(df.risk_factors)
    if factors:
        notes.append(f"Factors: {factors}")

    if df.fixed_version:
        notes.append(f"Fix: {df.fixed_version}")

    return " ".join(note for note in notes if note)


def _format_related_advisory(advisory: dict) -> str:
    advisory_id = (
        advisory.get("display_id")
        or advisory.get("cve_id")
        or advisory.get("advisory_id")
        or "advisory"
    )
    parts = [str(advisory_id)]

    severity = str(advisory.get("severity") or "").strip()
    if severity:
        parts.append(severity.upper())

    package = str(advisory.get("package") or "").strip()
    ecosystem = str(advisory.get("ecosystem") or "").strip()
    if package:
        package_label = package
        if ecosystem:
            package_label += f" ({ecosystem})"
        parts.append(package_label)

    evidence_strength = str(advisory.get("evidence_strength") or "").strip()
    evidence_type = str(advisory.get("evidence_type") or "").strip()
    evidence_bits = [bit for bit in [
        _humanise_dependency_label(evidence_strength),
        _humanise_dependency_label(evidence_type),
    ] if bit]
    if evidence_bits:
        parts.append(", ".join(evidence_bits))

    if advisory.get("import_module"):
        parts.append(f"import {advisory['import_module']}")
    if advisory.get("function"):
        function_label = f"call {advisory['function']}"
        if isinstance(advisory.get("line"), int):
            function_label += f" line {advisory['line']}"
        parts.append(function_label)
    if advisory.get("fixed_version"):
        parts.append(f"fix {advisory['fixed_version']}")

    summary = _truncate_text(str(advisory.get("summary") or "").strip(), 120)
    if summary:
        parts.append(summary)

    return " | ".join(part for part in parts if part)


def _format_related_advisories(advisories: list | None, limit: int = 3) -> list[str]:
    if not advisories:
        return []
    lines: list[str] = []
    for advisory in advisories[:limit]:
        if isinstance(advisory, dict):
            lines.append(_format_related_advisory(advisory))
    return [line for line in lines if line]


def _docx_diagram_image_bytes(diagram: RenderedReportDiagram) -> bytes:
    if not diagram.image_bytes:
        raise ValueError("No diagram bytes available")
    if is_svg_bytes(diagram.image_bytes):
        return svg_to_png_bytes(diagram.image_bytes)
    return diagram.image_bytes


async def _load_export_secrets(scan_id, db: AsyncSession) -> list[ExportSecretRow]:
    secrets_result = await db.execute(
        select(SecretCandidate).where(
            SecretCandidate.scan_id == scan_id,
            SecretCandidate.is_false_positive == False,
        )
    )
    secrets = secrets_result.scalars().all()

    file_ids = {secret.file_id for secret in secrets if secret.file_id}
    file_path_map: dict = {}
    if file_ids:
        file_rows = await db.execute(select(File.id, File.path).where(File.id.in_(file_ids)))
        file_path_map = {row.id: row.path for row in file_rows.all()}

    return [
        ExportSecretRow(
            type=secret.type,
            confidence=secret.confidence,
            file_path=file_path_map.get(secret.file_id) if secret.file_id else None,
        )
        for secret in secrets
    ]


async def generate_export(report: Report, format: str, db: AsyncSession) -> ExportArtifact:
    findings_result = await db.execute(
        select(Finding)
        .where(Finding.scan_id == report.scan_id)
        .options(selectinload(Finding.evidence))
        .order_by(Finding.severity)
    )
    findings = _sort_findings(findings_result.scalars().all())

    secrets = await _load_export_secrets(report.scan_id, db)

    dep_findings_result = await db.execute(
        select(DependencyFinding, Dependency)
        .join(Dependency, DependencyFinding.dependency_id == Dependency.id)
        .where(DependencyFinding.scan_id == report.scan_id)
    )
    dep_findings = _sort_dependency_findings(dep_findings_result.all())

    report_data = {
        "report": report,
        "findings": findings,
        "secrets": secrets,
        "dep_findings": dep_findings,
        "diagrams": await render_report_diagrams(report),
    }
    report_html = _render_report_html(report_data)
    report.report_html = report_html
    report_data["report_html"] = report_html

    settings.export_dir.mkdir(parents=True, exist_ok=True)
    file_name = f"vragent_report_{report.scan_id}.{format}"
    file_path = settings.export_dir / file_name

    if format == "docx":
        await _generate_docx(report_data, file_path)
    elif format == "pdf":
        await _generate_pdf(report_data, file_path)

    artifact = ExportArtifact(
        report_id=report.id,
        format=format,
        file_path=str(file_path),
        file_size=file_path.stat().st_size if file_path.exists() else 0,
    )
    db.add(artifact)
    await db.flush()
    return artifact


# ══════════════════════════════════════════════════════════════════
# DOCX GENERATION
# ══════════════════════════════════════════════════════════════════

async def _generate_docx(data: dict, output_path: Path):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    report: Report = data["report"]
    findings: list[Finding] = data["findings"]
    secrets = data["secrets"]
    dep_findings = data["dep_findings"]
    diagrams: list[RenderedReportDiagram] = data.get("diagrams", [])

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    section_num = 0
    def sec():
        nonlocal section_num
        section_num += 1
        return section_num

    # Split findings by severity tier
    critical_high = [f for f in findings if f.severity in ("critical", "high")]
    medium = [f for f in findings if f.severity == "medium"]
    low_info = [f for f in findings if f.severity in ("low", "info")]

    # ── Cover Page ─────────────────────────────────────────────
    doc.add_paragraph()  # spacing
    doc.add_paragraph()
    title = doc.add_heading("Security Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if report.app_summary:
        # Extract first line as app name
        first_line = report.app_summary.split("\n")[0].strip()
        if len(first_line) < 100:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub.add_run(first_line)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}\n").font.size = Pt(10)
    meta.add_run("VRAgent — AI-Assisted Vulnerability Research Platform").font.size = Pt(9)

    if report.risk_grade:
        doc.add_paragraph()
        grade_p = doc.add_paragraph()
        grade_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = grade_p.add_run(f"Risk Grade: {report.risk_grade}")
        run.bold = True
        run.font.size = Pt(24)
        grade_colors = {"A": 0x155724, "B": 0x0C5460, "C": 0x856404, "D": 0xE67E22, "F": 0xDC3545}
        run.font.color.rgb = RGBColor(*_int_to_rgb(grade_colors.get(report.risk_grade, 0x333333)))
        grade_p.add_run(f"  ({report.risk_score:.0f}/100)").font.size = Pt(14)

    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────
    n = sec()
    doc.add_heading(f"{n}. Executive Summary", level=1)

    grade_text = {
        "A": "No significant security issues identified. The application demonstrates strong security practices.",
        "B": "Minor issues found with low overall risk. Recommended to address findings in the next development cycle.",
        "C": "Moderate security concerns requiring attention. Several vulnerabilities should be addressed before production deployment.",
        "D": "Significant security issues discovered. Remediation is strongly recommended before production use.",
        "F": "Critical security failures identified. Immediate remediation is required. The application should not be deployed in its current state.",
    }
    if report.risk_grade:
        doc.add_paragraph(grade_text.get(report.risk_grade, ""))

    # Key metrics
    p = doc.add_paragraph()
    p.add_run("Key Metrics: ").bold = True
    p.add_run(
        f"{len(findings)} findings total — "
        f"{len([f for f in findings if f.severity == 'critical'])} Critical, "
        f"{len([f for f in findings if f.severity == 'high'])} High, "
        f"{len(medium)} Medium, "
        f"{len(low_info)} Low/Info"
    )
    exploitable = [f for f in findings if _finding_has_exploit_evidence(f)]
    if exploitable:
        p = doc.add_paragraph()
        p.add_run(f"Exploitable findings with PoC: ").bold = True
        p.add_run(str(len(exploitable)))

    advisory_count = len([f for f in findings if f.related_cves])
    if advisory_count:
        p = doc.add_paragraph()
        p.add_run(f"Findings with advisory correlations: ").bold = True
        p.add_run(str(advisory_count))

    # ── Application Overview ───────────────────────────────────
    n = sec()
    doc.add_heading(f"{n}. Application Overview", level=1)
    if report.app_summary:
        for para in _text_blocks(report.app_summary):
            doc.add_paragraph(para)

    if report.tech_stack:
        tech = report.tech_stack if isinstance(report.tech_stack, dict) else {}
        fp = tech.get("fingerprint", tech)
        langs = fp.get("languages", tech.get("languages", []))
        fws = fp.get("frameworks", tech.get("frameworks", []))
        if langs:
            items = langs if isinstance(langs[0], str) else [f"{l['name']} ({l['file_count']} files)" for l in langs]
            doc.add_paragraph(f"Languages: {', '.join(items)}")
        if fws:
            items = fws if isinstance(fws[0], str) else [f['name'] for f in fws]
            doc.add_paragraph(f"Frameworks: {', '.join(items)}")

    if report.narrative:
        n = sec()
        doc.add_heading(f"{n}. Security Review", level=1)
        for para in _text_blocks(report.narrative):
            doc.add_paragraph(para)

    # ── Architecture ───────────────────────────────────────────
    if diagrams:
        n = sec()
        doc.add_heading(f"{n}. Architecture", level=1)
        for index, diagram in enumerate(diagrams, start=1):
            if len(diagrams) > 1:
                doc.add_heading(f"{n}.{index}. {diagram.title}", level=2)
            elif diagram.title:
                doc.add_paragraph(diagram.title).runs[0].bold = True
            if diagram.description:
                doc.add_paragraph(diagram.description)
            for highlight in diagram.highlights[:4]:
                doc.add_paragraph(highlight, style="List Bullet")
            try:
                image_bytes = _docx_diagram_image_bytes(diagram)
                doc.add_picture(io.BytesIO(image_bytes), width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph("[Diagram could not be embedded]")

    attack_surface_points = _collect_attack_surface_points(report)
    trust_boundaries = _collect_trust_boundaries(report)
    if attack_surface_points or trust_boundaries:
        n = sec()
        doc.add_heading(f"{n}. Concrete Attack Surface", level=1)
        if attack_surface_points:
            doc.add_paragraph(
                "These concrete externally reachable routes, handlers, or invocations were identified during the architecture pass."
            )
            for point in attack_surface_points[:12]:
                doc.add_paragraph(point, style="List Bullet")
        if trust_boundaries:
            doc.add_paragraph("Trust boundaries observed:")
            for boundary in trust_boundaries[:8]:
                doc.add_paragraph(boundary, style="List Bullet")

    # ── Critical & High Findings (Full Detail) ─────────────────
    n = sec()
    doc.add_heading(f"{n}. Critical & High Severity Findings ({len(critical_high)})", level=1)

    if not critical_high:
        doc.add_paragraph("No critical or high severity findings were identified.")
    else:
        for i, f in enumerate(critical_high, 1):
            doc.add_heading(f"{n}.{i}. {f.title}", level=2)

            # Metadata line
            p = doc.add_paragraph()
            sev_run = p.add_run(f"{f.severity.upper()}")
            sev_run.bold = True
            sev_run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45) if f.severity == "critical" else RGBColor(0xE6, 0x7E, 0x22)
            p.add_run(f"  |  Confidence: {f.confidence:.0%}")
            if f.cwe_ids:
                p.add_run(f"  |  {', '.join(f.cwe_ids[:3])}")
            if f.category:
                p.add_run(f"  |  {f.category}")

            if f.description:
                doc.add_paragraph(f.description[:600])
            if f.explanation:
                p = doc.add_paragraph()
                p.add_run("Analysis: ").bold = True
                p.add_run(f.explanation[:900])

            if f.code_snippet:
                code_p = doc.add_paragraph()
                run = code_p.add_run(_truncate_code(f.code_snippet))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            if f.impact:
                p = doc.add_paragraph()
                p.add_run("Impact: ").bold = True
                p.add_run(f.impact[:300])

            if f.remediation:
                p = doc.add_paragraph()
                p.add_run("Remediation: ").bold = True
                p.add_run(f.remediation[:300])

            _add_docx_exploit_evidence(doc, f, heading_level=3)

            if f.related_cves:
                p = doc.add_paragraph()
                p.add_run("Related advisories: ").bold = True
                advisory_lines = _format_related_advisories(f.related_cves, limit=3)
                p.add_run(" ; ".join(advisory_lines))

    # ── Medium Findings (Concise) ──────────────────────────────
    if medium:
        n = sec()
        doc.add_heading(f"{n}. Medium Severity Findings ({len(medium)})", level=1)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "Title", "Category", "CWE", "Confidence"]):
            table.rows[0].cells[i].text = h

        for i, f in enumerate(medium, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = f.title[:80]
            row[2].text = f.category or ""
            row[3].text = ", ".join(f.cwe_ids[:2]) if f.cwe_ids else ""
            row[4].text = f"{f.confidence:.0%}"

        doc.add_paragraph()
        # Brief descriptions for medium findings (no code)
        for i, f in enumerate(medium, 1):
            if f.description:
                p = doc.add_paragraph()
                p.add_run(f"{i}. {f.title}: ").bold = True
                p.add_run(f.description[:200])

    appendix_findings = [
        f for f in findings
        if f.severity not in ("critical", "high") and _finding_has_exploit_evidence(f)
    ]
    if appendix_findings:
        n = sec()
        doc.add_heading(f"{n}. Exploit Evidence Appendix ({len(appendix_findings)})", level=1)
        doc.add_paragraph(
            "Structured exploit evidence is included here for findings outside the critical/high section "
            "so routes, validation steps, cleanup, and proof-of-concept details remain available in the export."
        )
        for i, f in enumerate(appendix_findings, 1):
            doc.add_heading(f"{n}.{i}. {f.title}", level=2)
            p = doc.add_paragraph()
            p.add_run(f"{f.severity.upper()}").bold = True
            p.add_run(f"  |  Confidence: {f.confidence:.0%}")
            if f.category:
                p.add_run(f"  |  {f.category}")
            if f.description:
                doc.add_paragraph(f.description[:400])
            _add_docx_exploit_evidence(doc, f, heading_level=3, heading_text=None)

    exploit_chains = [f for f in findings if str(f.category or "").lower() == "exploit_chain"]
    if exploit_chains:
        n = sec()
        doc.add_heading(f"{n}. Exploit Chains ({len(exploit_chains)})", level=1)
        doc.add_paragraph(
            "These findings describe multi-step attack paths where multiple weaknesses can be chained into a larger compromise."
        )
        for i, finding in enumerate(exploit_chains, 1):
            doc.add_heading(f"{n}.{i}. {finding.title}", level=2)
            p = doc.add_paragraph()
            p.add_run(f"{finding.severity.upper()}").bold = True
            p.add_run(f"  |  Confidence: {finding.confidence:.0%}")
            if finding.category:
                p.add_run(f"  |  {finding.category}")
            if finding.description:
                doc.add_paragraph(_truncate_text(finding.description, 700))
            if finding.impact:
                p = doc.add_paragraph()
                p.add_run("Impact: ").bold = True
                p.add_run(_truncate_text(finding.impact, 500))
            steps = _extract_exploit_chain_steps(finding)
            if steps:
                doc.add_paragraph("Chain steps:")
                for step in steps[:8]:
                    doc.add_paragraph(step, style="List Bullet")

    # ── Low & Info Findings (Table Only) ───────────────────────
    if low_info:
        n = sec()
        doc.add_heading(f"{n}. Low & Informational Findings ({len(low_info)})", level=1)
        doc.add_paragraph("These findings represent low-risk issues or informational observations.")

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Title", "Severity", "Category", "Confidence"]):
            table.rows[0].cells[i].text = h
        for f in low_info:
            row = table.add_row().cells
            row[0].text = f.title[:80]
            row[1].text = f.severity.upper()
            row[2].text = f.category or ""
            row[3].text = f"{f.confidence:.0%}"

    # ── OWASP Top 10 ──────────────────────────────────────────
    if report.owasp_mapping:
        n = sec()
        doc.add_heading(f"{n}. OWASP Top 10 Mapping", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Code", "Category", "Findings", "Max Severity"]):
            table.rows[0].cells[i].text = h
        for code in ["A01","A02","A03","A04","A05","A06","A07","A08","A09","A10"]:
            entry = report.owasp_mapping.get(code)
            if entry and entry.get("count", 0) > 0:
                row = table.add_row().cells
                row[0].text = code
                row[1].text = entry.get("name", "")
                row[2].text = str(entry["count"])
                row[3].text = entry.get("max_severity", "").upper()

    # ── Component Scorecard ────────────────────────────────────
    if report.component_scores:
        n = sec()
        doc.add_heading(f"{n}. Component Security Scorecard", level=1)
        doc.add_paragraph(
            "Components are graded A-F based on finding density and severity. "
            "A = no findings, B = low-severity only, C = medium findings present, "
            "D = high-severity findings, F = critical findings detected."
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Component", "Grade", "Score", "Criticality", "Findings"]):
            table.rows[0].cells[i].text = h
        for name, comp in sorted(report.component_scores.items(), key=lambda x: x[1]["score"]):
            row = table.add_row().cells
            row[0].text = name
            row[1].text = comp["grade"]
            row[2].text = str(comp["score"])
            row[3].text = comp.get("criticality", "")
            row[4].text = str(comp["finding_count"])

    if report.sbom and report.sbom.get("total_components", 0) > 0:
        sbom = report.sbom
        n = sec()
        doc.add_heading(f"{n}. Software Bill Of Materials", level=1)
        summary = doc.add_paragraph()
        summary.add_run("Component inventory: ").bold = True
        summary.add_run(
            f"{sbom.get('total_components', 0)} total components across "
            f"{len(sbom.get('ecosystems') or {})} ecosystems; "
            f"{sbom.get('vulnerable_components', 0)} marked vulnerable."
        )
        ecosystems = sbom.get("ecosystems") or {}
        if ecosystems:
            doc.add_paragraph(
                "Ecosystems: " + ", ".join(f"{name}={count}" for name, count in ecosystems.items())
            )

        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Component", "Version", "Ecosystem", "Scope", "Status", "Vulns"]):
            table.rows[0].cells[i].text = h

        components = sbom.get("components") or []
        for component in components[:30]:
            row = table.add_row().cells
            row[0].text = str(component.get("name") or "")[:80]
            row[1].text = str(component.get("version") or "—")[:40]
            row[2].text = str(component.get("ecosystem") or "")[:20]
            row[3].text = "dev" if component.get("is_dev") else "prod"
            row[4].text = "vulnerable" if component.get("vulnerable") else "ok"
            row[5].text = str(component.get("vulnerability_count") or 0)
        if len(components) > 30:
            doc.add_paragraph(f"... and {len(components) - 30} additional components.")

    # ── Secrets ────────────────────────────────────────────────
    if secrets:
        n = sec()
        doc.add_heading(f"{n}. Secrets & Sensitive Data ({len(secrets)})", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Type", "File", "Confidence"]):
            table.rows[0].cells[i].text = h
        for s in secrets[:30]:
            row = table.add_row().cells
            row[0].text = s.type
            row[1].text = (s.file_path or "")[:60]
            row[2].text = f"{s.confidence:.0%}" if s.confidence else ""
        if len(secrets) > 30:
            doc.add_paragraph(f"... and {len(secrets) - 30} additional secrets detected.")

    # ── Dependencies ───────────────────────────────────────────
    if dep_findings:
        n = sec()
        doc.add_heading(f"{n}. Dependency Risks ({len(dep_findings)})", level=1)
        reachable = sum(1 for df, _ in dep_findings if df.reachability_status == "reachable")
        active = sum(1 for df, _ in dep_findings if df.relevance in {"used", "likely_used"})
        doc.add_paragraph(
            "Dependency issues are ranked by combined risk score rather than advisory severity alone. "
            f"{reachable} are marked reachable and {active} show direct or likely package usage."
        )

        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Package", "Advisory", "Severity", "Exposure", "Risk", "Notes"]):
            table.rows[0].cells[i].text = h
        for df, dep in dep_findings[:20]:
            row = table.add_row().cells
            row[0].text = f"{dep.name}\n{dep.version or 'unknown'} ({dep.ecosystem})"
            row[1].text = df.cve_id or df.advisory_id or ""
            row[2].text = (df.severity or "").upper()
            row[3].text = _format_dependency_exposure(df)
            row[4].text = _format_dependency_risk_score(df.risk_score)
            row[5].text = _truncate_text(_format_dependency_notes(df) or "No package usage context captured.", 220)
        if len(dep_findings) > 20:
            doc.add_paragraph(f"... and {len(dep_findings) - 20} additional dependency risks.")

    # ── Scan Coverage ─────────────────────────────────────────
    if report.scan_coverage:
        cov = report.scan_coverage
        n = sec()
        doc.add_heading(f"{n}. Scan Coverage", level=1)
        total_files = int(cov.get("total_files", 0) or 0)
        ai_files = int(cov.get("files_inspected_by_ai", 0) or 0)
        ai_pct = round((ai_files / total_files) * 100) if total_files > 0 else 0
        p = doc.add_paragraph()
        p.add_run("Files indexed: ").bold = True
        p.add_run(str(cov.get("files_indexed", total_files)))
        p.add_run("  |  ")
        p.add_run("Total files: ").bold = True
        p.add_run(str(total_files))
        p.add_run("  |  ")
        p.add_run("AI inspected: ").bold = True
        p.add_run(f"{ai_files} ({ai_pct}%)")
        p.add_run("  |  ")
        p.add_run("AI calls made: ").bold = True
        p.add_run(str(cov.get("ai_calls_made", 0)))
        p.add_run("  |  ")
        p.add_run("Scan mode: ").bold = True
        p.add_run(str(cov.get("scan_mode", "?")))
        if cov.get("scanners_used"):
            p = doc.add_paragraph()
            p.add_run("Scanners: ").bold = True
            p.add_run(", ".join(cov["scanners_used"]))
        flags = _scan_coverage_flags(cov)
        if flags:
            doc.add_paragraph("Coverage notes:")
            for flag in flags:
                doc.add_paragraph(flag, style="List Bullet")
        scanner_runs = _scanner_run_rows(cov)
        if scanner_runs:
            doc.add_paragraph("Scanner run status:")
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(["Scanner", "Status", "Hits", "Errors"]):
                table.rows[0].cells[i].text = h
            for run in scanner_runs:
                row = table.add_row().cells
                row[0].text = str(run.get("scanner") or "")
                row[1].text = str(run.get("status") or "")
                row[2].text = str(run.get("hit_count") or 0)
                row[3].text = _truncate_text("; ".join(run.get("errors") or []), 140)
        availability = cov.get("scanner_availability") or {}
        if availability:
            doc.add_paragraph(
                "Scanner availability: "
                + ", ".join(f"{name}={status}" for name, status in sorted(availability.items()))
            )
        managed_paths = cov.get("managed_paths_ignored") or []
        if managed_paths:
            doc.add_paragraph(
                "Managed exclusions: " + ", ".join(str(path) for path in managed_paths[:6])
            )
        ignored_paths = cov.get("ignored_paths") or []
        if ignored_paths:
            doc.add_paragraph(
                "Ignored paths: " + ", ".join(str(path) for path in ignored_paths[:6])
            )
        if cov.get("repo_ignore_file"):
            doc.add_paragraph(f"Repo ignore file: {cov.get('repo_ignore_file')}")

    # ── Methodology ────────────────────────────────────────────
    n = sec()
    doc.add_heading(f"{n}. Methodology & Limitations", level=1)
    if report.methodology:
        doc.add_paragraph(report.methodology[:1500])
    if report.limitations:
        doc.add_heading("Limitations", level=2)
        doc.add_paragraph(report.limitations[:800])

    # ── Charts ─────────────────────────────────────────────────
    chart_images = _render_charts(report, findings, dep_findings)
    if chart_images:
        n = sec()
        doc.add_heading(f"{n}. Analytics", level=1)
        for chart_name, img_bytes in chart_images.items():
            try:
                doc.add_paragraph(chart_name)
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    doc.save(str(output_path))


def _int_to_rgb(val: int) -> tuple:
    return ((val >> 16) & 0xFF, (val >> 8) & 0xFF, val & 0xFF)


# ══════════════════════════════════════════════════════════════════
# PDF GENERATION (HTML → WeasyPrint)
# ══════════════════════════════════════════════════════════════════

async def _generate_pdf(data: dict, output_path: Path):
    import asyncio
    html = str(data.get("report_html") or _render_report_html(data))

    def _render():
        from weasyprint import HTML
        HTML(string=html).write_pdf(str(output_path))

    try:
        await asyncio.to_thread(_render)
    except Exception:
        await asyncio.to_thread(_generate_pdf_with_reportlab, data, output_path)


def _generate_pdf_with_reportlab(data: dict, output_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    report: Report = data["report"]
    findings: list[Finding] = data["findings"]
    secrets = data["secrets"]
    dep_findings = data["dep_findings"]

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "VRAgentTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#0a3d62"),
        spaceAfter=6,
    )
    section_style = ParagraphStyle(
        "VRAgentSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#1e6b8a"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "VRAgentBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "VRAgentMeta",
        parent=body_style,
        textColor=colors.HexColor("#5c677d"),
        alignment=TA_CENTER,
    )
    small_style = ParagraphStyle(
        "VRAgentSmall",
        parent=body_style,
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#5c677d"),
    )

    def p(text: str | None, style=body_style):
        if not text:
            return None
        return Paragraph(_esc(str(text)).replace("\n", "<br/>"), style)

    story = []
    story.append(Paragraph("Security Assessment Report", title_style))
    created = getattr(report, "created_at", None)
    if created:
        story.append(Paragraph(_esc(str(created)), meta_style))
    story.append(Spacer(1, 0.18 * inch))

    sev_counts: dict[str, int] = {}
    for finding in findings:
        sev_counts[finding.severity] = sev_counts.get(finding.severity, 0) + 1

    summary_rows = [
        ["Risk grade", str(report.risk_grade or "n/a"), "Risk score", str(report.risk_score or "n/a")],
        [
            "Findings",
            str(len(findings)),
            "Critical / High",
            str((sev_counts.get("critical", 0) + sev_counts.get("high", 0))),
        ],
    ]
    summary_table = Table(summary_rows, colWidths=[1.2 * inch, 1.2 * inch, 1.4 * inch, 1.4 * inch])
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f4f7fb")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1a1a2e")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d7deea")),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(summary_table)
    story.append(Spacer(1, 0.16 * inch))

    story.append(Paragraph("Application Summary", section_style))
    if report.app_summary:
        for block in _text_blocks(report.app_summary)[:8]:
            para = p(_truncate_text(block, 1800))
            if para:
                story.append(para)
    else:
        story.append(Paragraph("No application summary available.", small_style))

    if report.narrative:
        story.append(Paragraph("Security Review", section_style))
        for block in _text_blocks(report.narrative)[:8]:
            para = p(_truncate_text(block, 1800))
            if para:
                story.append(para)

    if findings:
        story.append(Paragraph(f"Findings ({len(findings)})", section_style))
        finding_rows = [["Severity", "Title", "Category", "Confidence"]]
        for finding in findings:
            finding_rows.append(
                [
                    (finding.severity or "").upper(),
                    _truncate_text(finding.title, 100),
                    _truncate_text(finding.category or "", 36),
                    f"{finding.confidence:.0%}",
                ]
            )
        finding_table = Table(
            finding_rows,
            repeatRows=1,
            colWidths=[0.8 * inch, 3.5 * inch, 1.35 * inch, 0.85 * inch],
        )
        finding_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0a3d62")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d7deea")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("PADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(finding_table)

    if secrets:
        story.append(Paragraph(f"Secrets &amp; Sensitive Data ({len(secrets)})", section_style))
        secret_rows = [["Type", "File", "Confidence"]]
        for secret in secrets[:40]:
            secret_rows.append(
                [
                    _truncate_text(secret.type, 28),
                    _truncate_text(secret.file_path or "", 72),
                    f"{secret.confidence:.0%}" if secret.confidence else "",
                ]
            )
        secret_table = Table(secret_rows, repeatRows=1, colWidths=[1.4 * inch, 3.9 * inch, 0.9 * inch])
        secret_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0a3d62")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d7deea")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("PADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(secret_table)

    if dep_findings:
        story.append(Paragraph(f"Dependency Risks ({len(dep_findings)})", section_style))
        dep_rows = [["Package", "Advisory", "Severity", "Risk", "Exposure"]]
        for dep_finding, dependency in dep_findings[:25]:
            dep_rows.append(
                [
                    _truncate_text(f"{dependency.name} {dependency.version or ''}".strip(), 40),
                    _truncate_text(dep_finding.cve_id or dep_finding.advisory_id or "", 22),
                    (dep_finding.severity or "").upper(),
                    _format_dependency_risk_score(dep_finding.risk_score),
                    _truncate_text(_format_dependency_exposure(dep_finding), 48),
                ]
            )
        dep_table = Table(
            dep_rows,
            repeatRows=1,
            colWidths=[1.8 * inch, 1.25 * inch, 0.8 * inch, 0.8 * inch, 2.05 * inch],
        )
        dep_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0a3d62")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d7deea")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("PADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(dep_table)

    if report.scan_coverage:
        cov = report.scan_coverage
        story.append(Paragraph("Scan Coverage", section_style))
        coverage_lines = [
            f"Files indexed: {cov.get('files_indexed', cov.get('total_files', 0))}",
            f"Total files: {cov.get('total_files', 0)}",
            f"AI inspected: {cov.get('files_inspected_by_ai', 0)}",
            f"AI calls: {cov.get('ai_calls_made', 0)}",
            f"Mode: {cov.get('scan_mode', 'unknown')}",
        ]
        story.append(Paragraph(" | ".join(_esc(line) for line in coverage_lines), small_style))
        flags = _scan_coverage_flags(cov)
        if flags:
            story.append(Paragraph("Coverage notes: " + _esc(", ".join(flags[:8])), small_style))

    story.append(Spacer(1, 0.18 * inch))
    story.append(Paragraph("Generated by VRAgent", meta_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        title="VRAgent Security Assessment Report",
    )
    doc.build(story)


def _render_report_html(data: dict) -> str:
    report: Report = data["report"]
    findings: list[Finding] = data["findings"]
    secrets = data["secrets"]
    dep_findings = data["dep_findings"]
    diagrams: list[RenderedReportDiagram] = data.get("diagrams", [])

    critical_high = [f for f in findings if f.severity in ("critical", "high")]
    medium = [f for f in findings if f.severity == "medium"]
    low_info = [f for f in findings if f.severity in ("low", "info")]

    parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>",
        "<style>",
        "@page { size: A4; margin: 2cm; @bottom-center { content: 'Page ' counter(page) ' of ' counter(pages); font-size: 8px; color: #999; } }",
        "body { font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif; color: #1a1a2e; font-size: 10.5px; line-height: 1.6; }",
        "h1 { color: #0a3d62; border-bottom: 3px solid #0a3d62; padding-bottom: 8px; font-size: 22px; margin-top: 30px; }",
        "h2 { color: #1e6b8a; margin-top: 28px; font-size: 16px; border-bottom: 1px solid #e0e0e0; padding-bottom: 4px; }",
        "h3 { color: #2d6a4f; font-size: 13px; margin-top: 20px; }",
        "table { width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 9.5px; page-break-inside: auto; }",
        "th, td { border: 1px solid #ddd; padding: 5px 8px; text-align: left; }",
        "th { background: #0a3d62; color: white; font-weight: 600; font-size: 9px; text-transform: uppercase; letter-spacing: 0.3px; }",
        "tr:nth-child(even) { background: #f8f9fa; }",
        "tr { page-break-inside: avoid; }",
        "pre, code { font-family: 'Consolas', 'Courier New', monospace; font-size: 8.5px; }",
        "pre { background: #f5f5f5; padding: 8px 12px; border-radius: 4px; border-left: 3px solid #0a3d62; overflow-x: auto; white-space: pre-wrap; word-break: break-word; }",
        ".cover { text-align: center; page-break-after: always; padding-top: 120px; }",
        ".cover h1 { border: none; font-size: 32px; color: #0a3d62; }",
        ".cover .subtitle { font-size: 16px; color: #555; margin: 10px 0 40px; }",
        ".cover .meta { font-size: 11px; color: #888; }",
        ".grade-box { display: inline-block; font-size: 36px; font-weight: 900; width: 60px; height: 60px; line-height: 60px; text-align: center; border-radius: 12px; margin: 20px 0; }",
        ".grade-A { background: #d4edda; color: #155724; } .grade-B { background: #d1ecf1; color: #0c5460; }",
        ".grade-C { background: #fff3cd; color: #856404; } .grade-D { background: #ffeeba; color: #856404; }",
        ".grade-F { background: #f8d7da; color: #721c24; }",
        ".sev-critical { color: #dc3545; font-weight: 700; } .sev-high { color: #e67e22; font-weight: 700; }",
        ".sev-medium { color: #f39c12; font-weight: 600; } .sev-low { color: #27ae60; } .sev-info { color: #3498db; }",
        ".finding-card { border: 1px solid #e0e0e0; border-radius: 6px; padding: 14px; margin: 12px 0; page-break-inside: avoid; background: #fafafa; }",
        ".finding-card h3 { margin-top: 0; border: none; padding: 0; }",
        ".meta-line { font-size: 9.5px; color: #555; margin: 4px 0 8px; }",
        ".poc-block { background: #fff0f0; border-left: 3px solid #dc3545; }",
        "img { max-width: 100%; height: auto; }",
        ".diagram-card { border: 1px solid #e0e0e0; border-radius: 8px; padding: 12px; margin: 14px 0; background: #fafafa; page-break-inside: avoid; }",
        ".diagram-card h3 { margin: 0 0 6px; color: #1e6b8a; }",
        ".diagram-card ul { margin: 6px 0 0 18px; padding: 0; }",
        ".diagram-card li { margin: 2px 0; }",
        ".page-break { page-break-before: always; }",
        "hr { border: none; border-top: 1px solid #e0e0e0; margin: 20px 0; }",
        "</style></head><body>",
    ]

    section = [0]
    def sec():
        section[0] += 1
        return section[0]

    # ── Cover Page ─────────────────────────────────────────────
    parts.append("<div class='cover'>")
    parts.append("<h1>Security Assessment Report</h1>")
    if report.app_summary:
        first_line = report.app_summary.split("\n")[0].strip()
        if len(first_line) < 100:
            parts.append(f"<div class='subtitle'>{_esc(first_line)}</div>")
    if report.risk_grade:
        parts.append(f"<div class='grade-box grade-{report.risk_grade}'>{report.risk_grade}</div>")
        parts.append(f"<div style='font-size:14px;'>Score: {report.risk_score:.0f}/100</div>")
    parts.append(f"<div class='meta'>Generated: {datetime.now().strftime('%d %B %Y')}<br>")
    parts.append("VRAgent — AI-Assisted Vulnerability Research Platform</div>")
    parts.append("</div>")

    # ── Executive Summary ──────────────────────────────────────
    n = sec()
    parts.append(f"<h1>{n}. Executive Summary</h1>")

    grade_text = {
        "A": "No significant security issues identified.",
        "B": "Minor issues found with low overall risk.",
        "C": "Moderate security concerns requiring attention.",
        "D": "Significant security issues. Remediation recommended before production.",
        "F": "Critical security failures. Immediate remediation required.",
    }
    if report.risk_grade:
        parts.append(f"<p>{grade_text.get(report.risk_grade, '')}</p>")

    parts.append(f"<p><strong>Total Findings:</strong> {len(findings)} — "
                 f"<span class='sev-critical'>{len([f for f in findings if f.severity == 'critical'])} Critical</span>, "
                 f"<span class='sev-high'>{len([f for f in findings if f.severity == 'high'])} High</span>, "
                 f"<span class='sev-medium'>{len(medium)} Medium</span>, "
                 f"{len(low_info)} Low/Info</p>")

    exploitable = [f for f in findings if _finding_has_exploit_evidence(f)]
    if exploitable:
        parts.append(f"<p><strong>Exploitable (PoC available):</strong> {len(exploitable)}</p>")

    # ── Application Overview ───────────────────────────────────
    n = sec()
    parts.append(f"<h1>{n}. Application Overview</h1>")
    if report.app_summary:
        for para in _text_blocks(report.app_summary):
            parts.append(f"<p>{_esc(para)}</p>")

    if report.narrative:
        n = sec()
        parts.append(f"<h1>{n}. Security Review</h1>")
        for para in _text_blocks(report.narrative):
            parts.append(f"<p>{_esc(para)}</p>")

    # ── Architecture ───────────────────────────────────────────
    if diagrams:
        n = sec()
        parts.append(f"<h1>{n}. Architecture</h1>")
        for diagram in diagrams:
            parts.append("<div class='diagram-card'>")
            if diagram.title:
                parts.append(f"<h3>{_esc(diagram.title)}</h3>")
            if diagram.description:
                parts.append(f"<p>{_esc(diagram.description)}</p>")
            uri = _diagram_to_data_uri(diagram)
            if uri:
                parts.append(f'<img src="{uri}" alt="{_esc(diagram.title or "Architecture diagram")}">')
            else:
                parts.append("<p><em>Diagram could not be rendered.</em></p>")
            if diagram.highlights:
                parts.append("<ul>")
                for highlight in diagram.highlights[:4]:
                    parts.append(f"<li>{_esc(highlight)}</li>")
                parts.append("</ul>")
            parts.append("</div>")

    attack_surface_points = _collect_attack_surface_points(report)
    trust_boundaries = _collect_trust_boundaries(report)
    if attack_surface_points or trust_boundaries:
        n = sec()
        parts.append(f"<h1>{n}. Concrete Attack Surface</h1>")
        if attack_surface_points:
            parts.append("<p><strong>Exposed entry points and attack paths:</strong></p><ul>")
            for point in attack_surface_points[:20]:
                parts.append(f"<li>{_esc(point)}</li>")
            parts.append("</ul>")
        if trust_boundaries:
            parts.append("<p><strong>Trust boundaries:</strong></p><ul>")
            for boundary in trust_boundaries[:10]:
                parts.append(f"<li>{_esc(boundary)}</li>")
            parts.append("</ul>")

    # ── Critical & High Findings ───────────────────────────────
    n = sec()
    parts.append(f"<h1>{n}. Critical & High Severity Findings ({len(critical_high)})</h1>")

    if not critical_high:
        parts.append("<p>No critical or high severity findings were identified.</p>")
    else:
        for i, f in enumerate(critical_high, 1):
            parts.append("<div class='finding-card'>")
            sev_class = f"sev-{f.severity}"
            parts.append(f"<h3>{n}.{i}. {_esc(f.title)}</h3>")
            meta = f"<span class='{sev_class}'>{f.severity.upper()}</span>"
            meta += f" &nbsp;|&nbsp; Confidence: {f.confidence:.0%}"
            if f.cwe_ids:
                meta += f" &nbsp;|&nbsp; {', '.join(f.cwe_ids[:3])}"
            if f.category:
                meta += f" &nbsp;|&nbsp; {_esc(f.category)}"
            parts.append(f"<div class='meta-line'>{meta}</div>")

            if f.description:
                parts.append(f"<p>{_esc(f.description[:600])}</p>")
            if f.explanation:
                parts.append(f"<p><strong>Analysis:</strong> {_esc(f.explanation[:900])}</p>")
            if f.code_snippet:
                parts.append(f"<pre><code>{_esc(_truncate_code(f.code_snippet))}</code></pre>")
            if f.impact:
                parts.append(f"<p><strong>Impact:</strong> {_esc(f.impact[:300])}</p>")
            if f.remediation:
                parts.append(f"<p><strong>Remediation:</strong> {_esc(f.remediation[:300])}</p>")
            parts.extend(_render_html_exploit_evidence(f))
            if f.related_cves:
                advisory_lines = ", ".join(_esc(line) for line in _format_related_advisories(f.related_cves, limit=3))
                parts.append(f"<p><strong>Related advisories:</strong> {advisory_lines}</p>")
            parts.append("</div>")

    # ── Medium Findings ────────────────────────────────────────
    if medium:
        n = sec()
        parts.append(f"<h1>{n}. Medium Severity Findings ({len(medium)})</h1>")
        parts.append("<table><tr><th>#</th><th>Title</th><th>Category</th><th>CWE</th><th>Confidence</th></tr>")
        for i, f in enumerate(medium, 1):
            cwe_str = ", ".join(f.cwe_ids[:2]) if f.cwe_ids else ""
            parts.append(f"<tr><td>{i}</td><td>{_esc(f.title[:80])}</td><td>{_esc(f.category or '')}</td><td>{cwe_str}</td><td>{f.confidence:.0%}</td></tr>")
        parts.append("</table>")

        for i, f in enumerate(medium, 1):
            if f.description:
                parts.append(f"<p><strong>{i}. {_esc(f.title[:60])}:</strong> {_esc(f.description[:200])}</p>")

    appendix_findings = [
        f for f in findings
        if f.severity not in ("critical", "high") and _finding_has_exploit_evidence(f)
    ]
    if appendix_findings:
        n = sec()
        parts.append(f"<h1>{n}. Exploit Evidence Appendix ({len(appendix_findings)})</h1>")
        parts.append(
            "<p>Structured exploit evidence is included here for lower-severity findings so route, validation, cleanup, "
            "and proof-of-concept details remain available in the exported report.</p>"
        )
        for i, f in enumerate(appendix_findings, 1):
            parts.append("<div class='finding-card'>")
            parts.append(f"<h3>{n}.{i}. {_esc(f.title)}</h3>")
            parts.append(
                f"<div class='meta-line'><span class='sev-{f.severity}'>{f.severity.upper()}</span>"
                f" &nbsp;|&nbsp; Confidence: {f.confidence:.0%}"
                + (f" &nbsp;|&nbsp; {_esc(f.category)}" if f.category else "")
                + "</div>"
            )
            if f.description:
                parts.append(f"<p>{_esc(f.description[:400])}</p>")
            parts.extend(_render_html_exploit_evidence(f, heading_text=None))
            parts.append("</div>")

    chain_findings = [f for f in findings if (f.category or "").lower() == "exploit_chain"]
    if chain_findings:
        n = sec()
        parts.append(f"<h1>{n}. Exploit Chains ({len(chain_findings)})</h1>")
        parts.append(
            "<p>These findings describe multi-step attack paths where individual weaknesses can be chained into a higher-impact compromise.</p>"
        )
        for i, finding in enumerate(chain_findings, 1):
            parts.append("<div class='finding-card'>")
            parts.append(f"<h3>{n}.{i}. {_esc(finding.title)}</h3>")
            parts.append(
                f"<div class='meta-line'><span class='sev-{finding.severity}'>{finding.severity.upper()}</span>"
                f" &nbsp;|&nbsp; Confidence: {finding.confidence:.0%}"
                + (f" &nbsp;|&nbsp; {_esc(finding.category)}" if finding.category else "")
                + "</div>"
            )
            if finding.description:
                parts.append(f"<p>{_esc(finding.description[:500])}</p>")
            if finding.impact:
                parts.append(f"<p><strong>Impact:</strong> {_esc(_truncate_text(finding.impact, 320))}</p>")
            steps = _extract_exploit_chain_steps(finding)
            if steps:
                parts.append("<p><strong>Chain steps:</strong></p><ol>")
                for step in steps[:8]:
                    parts.append(f"<li>{_esc(step)}</li>")
                parts.append("</ol>")
            parts.append("</div>")

    # ── Low & Info ─────────────────────────────────────────────
    if low_info:
        n = sec()
        parts.append(f"<h1>{n}. Low & Informational ({len(low_info)})</h1>")
        parts.append("<table><tr><th>Title</th><th>Severity</th><th>Category</th><th>Confidence</th></tr>")
        for f in low_info:
            parts.append(f"<tr><td>{_esc(f.title[:80])}</td><td class='sev-{f.severity}'>{f.severity.upper()}</td><td>{_esc(f.category or '')}</td><td>{f.confidence:.0%}</td></tr>")
        parts.append("</table>")

    # ── OWASP ──────────────────────────────────────────────────
    if report.owasp_mapping:
        n = sec()
        parts.append(f"<h1>{n}. OWASP Top 10 Mapping</h1>")
        parts.append("<table><tr><th>Code</th><th>Category</th><th>Count</th><th>Max Severity</th></tr>")
        for code in ["A01","A02","A03","A04","A05","A06","A07","A08","A09","A10"]:
            entry = report.owasp_mapping.get(code)
            if entry and entry.get("count", 0) > 0:
                parts.append(f"<tr><td><strong>{code}</strong></td><td>{_esc(entry['name'])}</td><td>{entry['count']}</td><td class='sev-{entry['max_severity']}'>{entry['max_severity'].upper()}</td></tr>")
        parts.append("</table>")

    # ── Component Scorecard ────────────────────────────────────
    if report.component_scores:
        n = sec()
        parts.append(f"<h1>{n}. Component Security Scorecard</h1>")
        parts.append("<p style='font-size:9px; color:#666;'>Grading: A = no findings, B = low-severity only, C = medium findings, D = high-severity, F = critical findings</p>")
        parts.append("<table><tr><th>Component</th><th>Grade</th><th>Score</th><th>Criticality</th><th>Findings</th></tr>")
        for name, comp in sorted(report.component_scores.items(), key=lambda x: x[1]["score"]):
            parts.append(f"<tr><td>{_esc(name)}</td><td><strong>{comp['grade']}</strong></td><td>{comp['score']}</td><td>{comp.get('criticality','')}</td><td>{comp['finding_count']}</td></tr>")
        parts.append("</table>")

    if report.sbom and report.sbom.get("total_components", 0) > 0:
        sbom = report.sbom
        n = sec()
        parts.append(f"<h1>{n}. Software Bill Of Materials</h1>")
        ecosystem_count = len(sbom.get("ecosystems") or {})
        parts.append(
            f"<p><strong>Component inventory:</strong> {sbom.get('total_components', 0)} total components across "
            f"{ecosystem_count} ecosystems; {sbom.get('vulnerable_components', 0)} marked vulnerable.</p>"
        )
        ecosystems = sbom.get("ecosystems") or {}
        if ecosystems:
            parts.append(
                f"<p><strong>Ecosystems:</strong> {_esc(', '.join(f'{name}={count}' for name, count in ecosystems.items()))}</p>"
            )
        components = sbom.get("components") or []
        if components:
            parts.append("<table><tr><th>Component</th><th>Version</th><th>Ecosystem</th><th>Scope</th><th>Status</th><th>Vulns</th></tr>")
            for component in components[:30]:
                parts.append(
                    "<tr>"
                    f"<td>{_esc(str(component.get('name') or '')[:80])}</td>"
                    f"<td>{_esc(str(component.get('version') or '—')[:40])}</td>"
                    f"<td>{_esc(str(component.get('ecosystem') or '')[:20])}</td>"
                    f"<td>{'dev' if component.get('is_dev') else 'prod'}</td>"
                    f"<td>{'vulnerable' if component.get('vulnerable') else 'ok'}</td>"
                    f"<td>{int(component.get('vulnerability_count') or 0)}</td>"
                    "</tr>"
                )
            parts.append("</table>")
            if len(components) > 30:
                parts.append(f"<p><em>... and {len(components) - 30} additional components</em></p>")

    # ── Secrets ─────────────────────────────────────────────────
    if secrets:
        n = sec()
        parts.append(f"<h1>{n}. Secrets & Sensitive Data ({len(secrets)})</h1>")
        parts.append("<table><tr><th>Type</th><th>File</th><th>Confidence</th></tr>")
        for s in secrets[:30]:
            conf = f"{s.confidence:.0%}" if s.confidence else ""
            parts.append(f"<tr><td>{_esc(s.type)}</td><td>{_esc((s.file_path or '')[:60])}</td><td>{conf}</td></tr>")
        parts.append("</table>")
        if len(secrets) > 30:
            parts.append(f"<p><em>... and {len(secrets) - 30} additional secrets</em></p>")

    # ── Dependencies ────────────────────────────────────────────
    if dep_findings:
        n = sec()
        parts.append(f"<h1>{n}. Dependency Risks ({len(dep_findings)})</h1>")
        reachable = sum(1 for df, _ in dep_findings if df.reachability_status == "reachable")
        active = sum(1 for df, _ in dep_findings if df.relevance in {"used", "likely_used"})
        parts.append(
            "<p>Dependency issues are ranked by combined risk score rather than advisory severity alone. "
            f"{reachable} are marked reachable and {active} show direct or likely package usage.</p>"
        )
        parts.append("<table><tr><th>Package</th><th>Advisory</th><th>Severity</th><th>Exposure</th><th>Risk</th><th>Notes</th></tr>")
        for df, dep in dep_findings[:20]:
            notes = _truncate_text(_format_dependency_notes(df) or "No package usage context captured.", 220)
            package = _esc(f"{dep.name} ({dep.version or 'unknown'} | {dep.ecosystem})")
            parts.append(
                f"<tr><td>{package}</td>"
                f"<td>{_esc(df.cve_id or df.advisory_id or '')}</td>"
                f"<td class='sev-{df.severity or 'info'}'>{(df.severity or '').upper()}</td>"
                f"<td>{_esc(_format_dependency_exposure(df))}</td>"
                f"<td>{_esc(_format_dependency_risk_score(df.risk_score))}</td>"
                f"<td>{_esc(notes)}</td></tr>"
            )
        parts.append("</table>")
        if len(dep_findings) > 20:
            parts.append(f"<p><em>... and {len(dep_findings) - 20} additional risks</em></p>")

    # ── Scan Coverage ──────────────────────────────────────────
    if report.scan_coverage:
        cov = report.scan_coverage
        n = sec()
        parts.append(f"<h1>{n}. Scan Coverage</h1>")
        total_files = int(cov.get("total_files", 0) or 0)
        ai_files = int(cov.get("files_inspected_by_ai", 0) or 0)
        ai_pct = round((ai_files / total_files) * 100) if total_files > 0 else 0
        parts.append(
            f"<p><strong>Files indexed:</strong> {cov.get('files_indexed', total_files)} &nbsp;|&nbsp; "
            f"<strong>Total files:</strong> {total_files} &nbsp;|&nbsp; "
            f"<strong>AI inspected:</strong> {ai_files} ({ai_pct}%) &nbsp;|&nbsp; "
            f"<strong>AI calls:</strong> {cov.get('ai_calls_made', 0)} &nbsp;|&nbsp; "
            f"<strong>Mode:</strong> {_esc(str(cov.get('scan_mode', '?')))}</p>"
        )
        if cov.get("scanners_used"):
            parts.append(f"<p><strong>Scanners:</strong> {_esc(', '.join(str(s) for s in cov['scanners_used']))}</p>")
        flags = _scan_coverage_flags(cov)
        if flags:
            parts.append("<p><strong>Coverage notes:</strong></p><ul>")
            for flag in flags:
                parts.append(f"<li>{_esc(flag)}</li>")
            parts.append("</ul>")
        scanner_runs = _scanner_run_rows(cov)
        if scanner_runs:
            parts.append("<p><strong>Scanner run status:</strong></p>")
            parts.append("<table><tr><th>Scanner</th><th>Status</th><th>Hits</th><th>Errors</th></tr>")
            for run in scanner_runs:
                parts.append(
                    "<tr>"
                    f"<td>{_esc(str(run.get('scanner') or ''))}</td>"
                    f"<td>{_esc(str(run.get('status') or ''))}</td>"
                    f"<td>{int(run.get('hit_count') or 0)}</td>"
                    f"<td>{_esc(_truncate_text('; '.join(run.get('errors') or []), 140))}</td>"
                    "</tr>"
                )
            parts.append("</table>")
        availability = cov.get("scanner_availability") or {}
        if availability:
            parts.append(
                f"<p><strong>Scanner availability:</strong> {_esc(', '.join(f'{name}={status}' for name, status in sorted(availability.items())))}</p>"
            )
        managed_paths = cov.get("managed_paths_ignored") or []
        if managed_paths:
            parts.append(
                f"<p><strong>Managed exclusions:</strong> {_esc(', '.join(str(path) for path in managed_paths[:6]))}</p>"
            )
        ignored_paths = cov.get("ignored_paths") or []
        if ignored_paths:
            parts.append(
                f"<p><strong>Ignored paths:</strong> {_esc(', '.join(str(path) for path in ignored_paths[:6]))}</p>"
            )
        if cov.get("repo_ignore_file"):
            parts.append(f"<p><strong>Repo ignore file:</strong> {_esc(str(cov.get('repo_ignore_file')))}</p>")

    # ── Methodology ────────────────────────────────────────────
    n = sec()
    parts.append(f"<h1>{n}. Methodology & Limitations</h1>")
    if report.methodology:
        parts.append(f"<p>{_esc(report.methodology[:1500])}</p>")
    if report.limitations:
        parts.append(f"<p><strong>Limitations:</strong> {_esc(report.limitations[:800])}</p>")

    # ── Charts ─────────────────────────────────────────────────
    chart_images = _render_charts(report, findings, dep_findings)
    if chart_images:
        import base64 as _b64
        n = sec()
        parts.append(f"<h1 class='page-break'>{n}. Analytics</h1>")
        parts.append("<div style='display:flex; flex-wrap:wrap; gap:20px; justify-content:center;'>")
        for chart_name, img_bytes in chart_images.items():
            b64 = _b64.b64encode(img_bytes).decode("ascii")
            parts.append(f"<div style='text-align:center;'><img src='data:image/png;base64,{b64}' style='max-width:320px;'><p style='font-size:9px; color:#666;'>{_esc(chart_name)}</p></div>")
        parts.append("</div>")

    parts.append("<hr><p style='text-align:center; color:#aaa; font-size:8px;'>Generated by VRAgent — AI-Assisted Vulnerability Research Platform</p>")
    parts.append("</body></html>")
    return "\n".join(parts)


def _esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


# ══════════════════════════════════════════════════════════════════
# CHART RENDERING (matplotlib)
# ══════════════════════════════════════════════════════════════════

SEV_COLORS = {"critical": "#ef4444", "high": "#f97316", "medium": "#eab308", "low": "#22c55e", "info": "#06b6d4"}
SCANNER_COLORS = {"semgrep": "#06b6d4", "bandit": "#f59e0b", "eslint": "#8b5cf6", "codeql": "#ec4899", "secrets": "#ef4444", "dep_audit": "#22c55e"}


def _render_charts(
    report: Report,
    findings: list,
    dep_findings: list[tuple[DependencyFinding, Dependency]] | None = None,
) -> dict[str, bytes]:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return {}

    charts: dict[str, bytes] = {}
    plt.rcParams.update({"figure.facecolor": "white", "axes.facecolor": "white", "font.size": 9, "font.family": "sans-serif"})
    dep_findings = dep_findings or []

    def _save(fig, name: str):
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        charts[name] = buf.getvalue()

    # Export-safe subset aligned with the interactive dashboard: severity, scanners,
    # categories, confidence, attack surface, language footprint, and dependency risk.

    # 1. Severity donut
    if findings:
        sev_counts: dict[str, int] = {}
        for f in findings:
            sev_counts[f.severity] = sev_counts.get(f.severity, 0) + 1
        if sev_counts:
            fig, ax = plt.subplots(figsize=(4, 3))
            labels = list(sev_counts.keys())
            sizes = list(sev_counts.values())
            colors = [SEV_COLORS.get(s, "#999") for s in labels]
            ax.pie(sizes, labels=[s.title() for s in labels], colors=colors, autopct="%1.0f%%", startangle=90, pctdistance=0.8, wedgeprops={"width": 0.4, "edgecolor": "white", "linewidth": 1.5})
            ax.set_title("Severity Breakdown", fontsize=11, fontweight="bold", pad=12)
            _save(fig, "Severity Breakdown")

    # 2. Scanner hits
    if report.scanner_hits:
        hits = {k: v for k, v in report.scanner_hits.items() if v > 0}
        if hits:
            fig, ax = plt.subplots(figsize=(4.5, 2.5))
            names = list(hits.keys())
            values = list(hits.values())
            colors = [SCANNER_COLORS.get(n, "#666") for n in names]
            bars = ax.barh([n.replace("_", " ").title() for n in names], values, color=colors, edgecolor="white", linewidth=0.5, height=0.6)
            ax.set_title("Scanner Hit Distribution", fontsize=11, fontweight="bold")
            ax.set_xlabel("Hits")
            ax.spines["top"].set_visible(False)
            ax.spines["right"].set_visible(False)
            for bar, val in zip(bars, values):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2, str(val), va="center", fontsize=8)
            _save(fig, "Scanner Hits")

    # 3. Finding categories
    if findings:
        category_counts: dict[str, int] = {}
        for finding in findings:
            category = (finding.category or "uncategorized").strip().lower()
            category_counts[category] = category_counts.get(category, 0) + 1
        category_rows = sorted(category_counts.items(), key=lambda item: (-item[1], item[0]))[:10]
        if category_rows:
            fig, ax = plt.subplots(figsize=(5.2, 3.0))
            labels = [label.replace("_", " ").title() for label, _ in category_rows]
            values = [value for _, value in category_rows]
            bars = ax.bar(labels, values, color="#1d4ed8", alpha=0.85)
            ax.set_title("Finding Categories", fontsize=11, fontweight="bold")
            ax.set_ylabel("Findings")
            ax.tick_params(axis="x", labelrotation=35, labelsize=8)
            ax.spines["top"].set_visible(False)
            ax.spines["right"].set_visible(False)
            for bar, value in zip(bars, values):
                ax.text(bar.get_x() + bar.get_width() / 2, value + 0.1, str(value), ha="center", va="bottom", fontsize=8)
            _save(fig, "Finding Categories")

    # 4. Confidence distribution
    if findings:
        buckets = {"90-100%": 0, "70-89%": 0, "50-69%": 0, "<50%": 0}
        for f in findings:
            pct = f.confidence * 100
            if pct >= 90: buckets["90-100%"] += 1
            elif pct >= 70: buckets["70-89%"] += 1
            elif pct >= 50: buckets["50-69%"] += 1
            else: buckets["<50%"] += 1
        nonzero = {k: v for k, v in buckets.items() if v > 0}
        if nonzero:
            fig, ax = plt.subplots(figsize=(3.5, 3))
            ax.pie(list(nonzero.values()), labels=list(nonzero.keys()), colors=["#ef4444","#f97316","#eab308","#22c55e"][:len(nonzero)], autopct="%1.0f%%", startangle=90, wedgeprops={"edgecolor": "white", "linewidth": 1.5})
            ax.set_title("Confidence Distribution", fontsize=11, fontweight="bold", pad=12)
            _save(fig, "Confidence Distribution")

    # 5. Attack surface radar
    if report.attack_surface and len(report.attack_surface) >= 3:
        import numpy as np
        categories = list(report.attack_surface.keys())
        values = list(report.attack_surface.values())
        max_val = max(values) if values else 1
        normalized = [v / max_val * 100 for v in values]
        N = len(categories)
        angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
        normalized += normalized[:1]
        angles += angles[:1]
        fig, ax = plt.subplots(figsize=(4, 4), subplot_kw=dict(polar=True))
        ax.fill(angles, normalized, alpha=0.15, color="#06b6d4")
        ax.plot(angles, normalized, color="#06b6d4", linewidth=2)
        ax.scatter(angles[:-1], normalized[:-1], color="#06b6d4", s=30, zorder=5)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(categories, fontsize=7)
        ax.set_yticklabels([])
        ax.set_title("Attack Surface", fontsize=11, fontweight="bold", pad=20)
        _save(fig, "Attack Surface")

    # 6. Language distribution
    tech = report.tech_stack or {}
    fingerprint = tech.get("fingerprint") if isinstance(tech, dict) else {}
    language_rows = []
    if isinstance(fingerprint, dict):
        for language in fingerprint.get("languages") or []:
            if isinstance(language, dict):
                name = str(language.get("name") or "").strip()
                file_count = int(language.get("file_count") or 0)
                if name and file_count > 0:
                    language_rows.append((name, file_count))
    if language_rows:
        language_rows = sorted(language_rows, key=lambda item: (-item[1], item[0].lower()))[:8]
        fig, ax = plt.subplots(figsize=(4.8, 3.0))
        labels = [name for name, _ in language_rows]
        values = [count for _, count in language_rows]
        bars = ax.barh(labels, values, color="#0ea5e9", alpha=0.85)
        ax.set_title("Language Footprint", fontsize=11, fontweight="bold")
        ax.set_xlabel("Files")
        ax.invert_yaxis()
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        for bar, value in zip(bars, values):
            ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2, str(value), va="center", fontsize=8)
        _save(fig, "Language Footprint")

    # 7. Dependency risk
    if dep_findings:
        reachability_counts: dict[str, int] = {}
        severity_counts: dict[str, int] = {}
        has_reachability = False
        for dependency_finding, _dependency in dep_findings:
            status = str(dependency_finding.reachability_status or "").strip().lower()
            severity = str(dependency_finding.severity or "unknown").strip().lower()
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
            if status and status != "unknown":
                has_reachability = True
                reachability_counts[status] = reachability_counts.get(status, 0) + 1

        if has_reachability and reachability_counts:
            label_map = {
                "reachable": "Reachable",
                "potentially_reachable": "Potential",
                "no_path_found": "No path",
                "not_applicable": "Not applicable",
                "unknown": "Unknown",
            }
            color_map = {
                "reachable": "#ef4444",
                "potentially_reachable": "#f97316",
                "no_path_found": "#06b6d4",
                "not_applicable": "#22c55e",
                "unknown": "#6b7280",
            }
            rows = [(status, count) for status, count in reachability_counts.items() if count > 0]
            fig, ax = plt.subplots(figsize=(4, 3))
            ax.pie(
                [count for _, count in rows],
                labels=[label_map.get(status, status.title()) for status, _ in rows],
                colors=[color_map.get(status, "#6b7280") for status, _ in rows],
                autopct="%1.0f%%",
                startangle=90,
                pctdistance=0.8,
                wedgeprops={"width": 0.4, "edgecolor": "white", "linewidth": 1.5},
            )
            ax.set_title("Dependency Reachability", fontsize=11, fontweight="bold", pad=12)
            _save(fig, "Dependency Reachability")
        elif severity_counts:
            rows = [(severity, count) for severity, count in severity_counts.items() if count > 0]
            fig, ax = plt.subplots(figsize=(4, 3))
            ax.pie(
                [count for _, count in rows],
                labels=[severity.title() for severity, _ in rows],
                colors=[SEV_COLORS.get(severity, "#6b7280") for severity, _ in rows],
                autopct="%1.0f%%",
                startangle=90,
                pctdistance=0.8,
                wedgeprops={"width": 0.4, "edgecolor": "white", "linewidth": 1.5},
            )
            ax.set_title("Dependency Severity", fontsize=11, fontweight="bold", pad=12)
            _save(fig, "Dependency Severity")

    return charts
